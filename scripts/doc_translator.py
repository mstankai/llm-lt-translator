from openai import OpenAI
from docx import Document
from docx.shared import Pt
from docx.shared import RGBColor 
import tiktoken
import yaml
import keyring
import argparse
import sys
import os


def system_prompt(language: str) -> str:
    return f"""
    You are a skilled {language} language translator. Your goal is to provide translations that are:
    Accurate: Convey the exact meaning of the original text, including idioms and cultural references.
    Fluent: Ensure the translation reads naturally and smoothly in the target language.
    Culturally Appropriate: Adapt the text for the target audience while respecting cultural differences.
    Consistent: Maintain tone, style, and terminology throughout the translation.
    Context-Aware: Understand and accurately translate technical or specialized terms relevant to the text's field.
    This is really important: Make sure to keep the delimiter {DELIMITER} in the same places as in the original text.
    Make sure there is the same number of delimiters in the output as in the input.
    Please translate the following text from English to {language}, keeping the delimiter {DELIMITER} between the text blocks:
    """

def output_exists(output_path: str) -> None:
    # check if output file exists, if it does, ask to confirm overwrite
    if os.path.exists(output_path):
        print(f"\nOutput file {output_path} already exists." )
        print("Overwrite? (y / any key)")
        response = input()
        if response.lower() != 'y':
            print("Exiting...")
            sys.exit(0)


def get_api_key(key_name: str) -> str:
    api_key = keyring.get_password("system", key_name)
    if api_key is None:
        raise ValueError("API key not found in Keychain!")
    return api_key


def read_docx(file_path: str) -> tuple:
    
    doc = Document(file_path)

    texts = []
    formats = []

    for paragraph in doc.paragraphs:
        for run in paragraph.runs:

            run_format = {
                'bold': run.bold,
                'italic': run.italic,
                'underline': run.underline,
                'font_name': run.font.name,
                'font_size': run.font.size.pt if run.font.size else None,
                'font_color': run.font.color.rgb if run.font.color else None,
            }

            texts.append(run.text)
            formats.append(run_format)

    return texts, formats

def write_docx(output_path:str, texts: list, formats: list) -> None:
    
    doc = Document()
    
    for text, f in zip(texts, formats):
        run = doc.add_paragraph().add_run(text)
        run.bold = f['bold']
        run.italic = f['italic']
        run.underline = f['underline']
        run.font.name = f['font_name']
        
        if f['font_size'] is not None:
            run.font.size = Pt(f['font_size'])
        
        if f['font_color'] is not None:
            r, g, b = f['font_color']
            run.font.color.rgb = RGBColor(r, g, b)

    doc.save(output_path)
    print(f"Output saved to {output_path}")


def get_number_of_tokens(messages: list, model_name: str) -> int:

    # get tokenizer encoding
    encoding = tiktoken.encoding_for_model(model_name)

    # counting tokens from each message
    token_count = 0
    for m in messages:
        text = m['content']
        tokens  = encoding.encode(text)
        token_count += len(tokens)

    return token_count


def get_cost_of_tokens(
        n_tokens: int,
        model_name: str, 
        price_list: dict, 
        is_input: bool = True
    ) -> float:

    # get pice list keys for model
    model_group = MODEL_TYPE + "_models"
    token_type = "input_tokens" if is_input else "output_tokens"

    cost_per_1M_tokens = price_list[model_group][model_name][token_type]
    cost = cost_per_1M_tokens * n_tokens / 1e6

    return cost    

def get_completion(client: OpenAI, messages: list, model_name: str) -> str:
    
    completion_input = dict(
        model=model_name,
        messages=messages
    )
        
    completion = client.chat.completions.create(**completion_input)
    return completion.choices[0]


def compare_text_blocks(input_text_list: list, output_text_list: list) -> None:
    l_in = len(input_text_list)
    l_out = len(output_text_list)

    if l_in != l_out:
        print(
            f"WARNING! Output text block count ({l_out}) "
            f"does not match input text bolock count ({l_in})."
        )
        print("Formatting will not be reliable!\n")

        print("Do you want to compare the text blocks? (y / any key)")
        response = input()
        if response.lower() == 'y':
            print("\n-----------------------------------")
            print("Input and output text blocks:\n")
            print_text_block_comparison(input_text_list, output_text_list)
            print("\n-----------------------------------")

        
        print("Continuing...")


def print_text_block_comparison(input_text_list: list, output_text_list: list) -> None:
    l_in = len(input_text_list)
    l_out = len(output_text_list)
    i_max = max(l_in, l_out)

    for i in range(i_max):
        i_text = input_text_list[i] if i < l_in else "---"
        o_text = output_text_list[i] if i < l_out else "---"
        
        print(f" >> Input: {i_text}")
        print(f" >> Output: {o_text}")
        print()


# ------------------------------------------
def main() -> None:

    # get price list and available models
    with open(TOKEN_PRICE_LIST_PATH, 'r') as f:
        price_list = yaml.safe_load(f)

    available_models = price_list[MODEL_TYPE + "_models"].keys()

    # get arguments
    parser = argparse.ArgumentParser(description='Translate a PDF file to a selected language')
    parser.add_argument('input')
    parser.add_argument('-o','--output', default=None)
    parser.add_argument('-m','--model', default='gpt-4o-mini', choices=available_models)
    parser.add_argument('-l','--lang', default='Lithuanian')
    parser.add_argument('-v','--verbose', action='store_true')
    parser.add_argument('--dry-run', action='store_true')
        
    args = parser.parse_args()

    input_file = args.input
    output_file = args.output
    model_name = args.model
    language = args.lang
    verbose = args.verbose
    dry_run = args.dry_run

    # get output file name if not provided
    if output_file is None:
        input_file_name = os.path.basename(input_file).split(".")[0]
        output_file_name = input_file_name + "_translated.docx"
        output_file = f"./output/{output_file_name}"

    print(f"\nInput file: {input_file}")
    print(f"Output file: {output_file}")
    output_exists(output_file)

    # get text blocks and their formatting
    text_list, formatting_list = read_docx(input_file)
    
    # combine text into single block
    full_text = DELIMITER.join(text_list)

    # transform text to OpenAI messages format
    messages = [
        {'role': 'system', 'content': system_prompt(language)},
        {'role': 'user', 'content': full_text}
    ]

    # calculate tokens and est. cost
    n_input_tokens = get_number_of_tokens(messages, model_name)
    input_cost = get_cost_of_tokens(n_input_tokens, model_name, price_list, is_input=True)


    print(f"\nSubmitting {len(messages)} messages to OpenAI API...")
    print(f"Number of tokens: {n_input_tokens}")
    print(f"Estimated cost of input: {input_cost:3f} USD\n")

    if n_input_tokens > MAX_TOKENS:
        ValueError(f"Input tokens ({n_input_tokens}) exceed the maximum of {MAX_TOKENS}.")

    if verbose:
        print("\n-----------------------------------")
        print("Messages:\n")
        for m in messages:
            print(f" >> {m['role']}: {m['content']}\n")

    if dry_run:
        print("Dry run. Not submittion is made.\n")
        return
        
    # set up OpenAI API
    openai_api_key = get_api_key("openai_api_key")
    client = OpenAI(api_key=openai_api_key)

    # get completion
    completion=get_completion(client, messages, model_name)
    
    if verbose:
        print(completion, "\n")
    
    output_message = completion.message.content
    # output_message = TEST_OUTPUT

    output_text_list = output_message.split(DELIMITER)


    # check if output is the same length as input
    compare_text_blocks(text_list, output_text_list) 


    if verbose:
        print("\n-----------------------------------")
        print("Text blocks and their formatting:\n")
        for o, f in zip(output_text_list, formatting_list):
            print(f)
            print(o)
            print()
        print("-----------------------------------")

    # write output to docx
    write_docx(output_file, output_text_list, formatting_list)

# ------------------------------------------
if __name__ == "__main__":

    # constants
    TOKEN_PRICE_LIST_PATH = "etc/openai_pricing_23Oct2024.yaml"
    MODEL_TYPE = "chat"
    DELIMITER = " &#&#&#& "
    MAX_TOKENS = 5000

    TEST_INPUT = "Fiber Optics &#&#&#& Fiber optic systems consist of three main components: the optical fiber, a light source (transmitter), and a detector (receiver). &#&#&#& Hello! Is this printing?"
    TEST_OUTPUT = "Pluoštinė optika &#&#&#& Pluoštinės optikos sistemos susideda iš trijų pagrindinių komponentų: optinio pluošto, šviesos šaltinio (siuntėjo) ir detektoriaus (imtuvo). &#&#&#& Sveiki! Ar tai spausdinimas?"
    

    main()
    print("\nTranslation completed.")
